
import os
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from ..db.models import Analysis
from ..utils.logger import get_logger

logger = get_logger(__name__)

class ExportService:
    @staticmethod
    async def generate_pdf(analysis_id: str):
        """Generate a PDF file from an analysis"""
        try:
            # Get analysis from DB
            analysis = await Analysis.get(analysis_id)
            if not analysis:
                return None
            
            # Create PDF
            pdf = FPDF()
            pdf.add_page()
            
            # Set font
            pdf.set_font("Arial", "B", 16)
            
            # Title
            pdf.cell(200, 10, "Multi-Perspective Analysis", ln=True, align="C")
            pdf.set_font("Arial", "B", 12)
            pdf.cell(200, 10, f"Article: {analysis.article_title}", ln=True)
            pdf.cell(200, 10, f"Date: {datetime.now().strftime('%Y-%m-%d')}", ln=True)
            
            # Content
            pdf.set_font("Arial", "", 12)
            
            # Summary
            pdf.set_font("Arial", "B", 12)
            pdf.cell(200, 10, "Summary:", ln=True)
            pdf.set_font("Arial", "", 12)
            pdf.multi_cell(0, 10, analysis.summary)
            
            # Perspectives
            pdf.set_font("Arial", "B", 12)
            pdf.cell(200, 10, "Perspectives:", ln=True)
            
            for perspective in analysis.perspectives:
                pdf.set_font("Arial", "B", 11)
                pdf.cell(200, 10, f"- {perspective.name}:", ln=True)
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(0, 10, perspective.content)
            
            # Save to temp file
            filename = f"analysis_{analysis_id}.pdf"
            filepath = os.path.join("/tmp", filename)
            pdf.output(filepath)
            
            return filepath
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return None
    
    @staticmethod
    async def generate_docx(analysis_id: str):
        """Generate a DOCX file from an analysis"""
        try:
            # Get analysis from DB
            analysis = await Analysis.get(analysis_id)
            if not analysis:
                return None
            
            # Create DOCX
            doc = Document()
            
            # Title
            doc.add_heading("Multi-Perspective Analysis", 0)
            doc.add_heading(f"Article: {analysis.article_title}", 1)
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            
            # Summary
            doc.add_heading("Summary:", 2)
            doc.add_paragraph(analysis.summary)
            
            # Perspectives
            doc.add_heading("Perspectives:", 2)
            
            for perspective in analysis.perspectives:
                p = doc.add_paragraph()
                p.add_run(f"{perspective.name}:").bold = True
                doc.add_paragraph(perspective.content)
            
            # Save to temp file
            filename = f"analysis_{analysis_id}.docx"
            filepath = os.path.join("/tmp", filename)
            doc.save(filepath)
            
            return filepath
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return None